import requests
import re
from PIL import Image, ImageOps
import matplotlib.pyplot as plt
from collections import Counter
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# === CONFIGURATION ===
BOOK_URL = "https://www.gutenberg.org/files/11/11-0.txt"
IMAGE1_URL = "https://www.gutenberg.org/files/11/11-h/images/cover.jpg"
LOGO_PATH = "logo.png"
OUTPUT_DIR = "."
GRAPH_PATH = "graphique.png"
IMAGE1_PATH = "alice_image.jpg"
MERGED_IMAGE_PATH = "image_finale.jpg"
DOCX_PATH = "rapport_final.docx"
AUTEUR_RAPPORT = "Kervoelen Erwann"  # Valeur par défaut

def download_book():
    response = requests.get(BOOK_URL)
    response.encoding = 'utf-8'
    return response.text

def extract_info(text):
    text = text.replace('\r\n', '\n')
    lines = text.splitlines()

    title = "Titre inconnu"
    for line in lines:
        clean = line.strip()
        if (10 < len(clean) < 100 and not any(word in clean.upper() for word in ["PROJECT", "GUTENBERG", "ILLUSTRATION", "***"])):
            title = clean
            break

    author_line = next((line.strip() for line in lines if line.strip().lower().startswith("by ")), None)
    author = author_line[3:].strip() if author_line else "Auteur inconnu"

    print(f"\n✔️ Titre : {title}")
    print(f"✔️ Auteur : {author}")

    match = re.search(r"CHAPTER I\..*?\n(.*?)\nCHAPTER II\.", text, re.DOTALL | re.IGNORECASE)
    if not match:
        raise ValueError("Impossible d’extraire le contenu entre CHAPTER I et II.")

    first_chapter = match.group(1).strip()
    print(f"Taille brute du chapitre récupéré : {len(first_chapter)} caractères")

    lines = first_chapter.splitlines()
    paragraphs = []
    current_paragraph = []

    for line in lines:
        stripped = line.strip()
        if stripped == "":
            if current_paragraph:
                joined = " ".join(current_paragraph).strip()
                if len(joined.split()) > 10:
                    paragraphs.append(joined)
                current_paragraph = []
        else:
            current_paragraph.append(stripped)

    if current_paragraph:
        joined = " ".join(current_paragraph).strip()
        if len(joined.split()) > 10:
            paragraphs.append(joined)

    if not paragraphs:
        raise ValueError("Aucun paragraphe extrait. Vérifiez le format du texte.")

    print(f"✅ {len(paragraphs)} paragraphes extraits.")
    print("Exemple :", paragraphs[0][:100], "...")
    return title, author, paragraphs

def analyze_paragraphs(paragraphs):
    lengths = [len(p.split()) for p in paragraphs]
    rounded = [int(round(l, -1)) for l in lengths]
    counter = Counter(rounded)
    sorted_items = sorted(counter.items())

    x, y = zip(*sorted_items)
    plt.figure(figsize=(6, 4))
    plt.bar(x, y)
    plt.xlabel("Nombre de mots (arrondi à la dizaine)")
    plt.ylabel("Nombre de paragraphes")
    plt.title("Distribution des longueurs de paragraphes")
    plt.tight_layout()
    plt.savefig(GRAPH_PATH)
    plt.close()
    return lengths

def download_image():
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) PythonBot/1.0"
    }
    response = requests.get(IMAGE1_URL, headers=headers)

    if response.status_code != 200 or b"<html" in response.content[:100].lower():
        raise ValueError("L'image n'a pas été correctement téléchargée.")

    with open(IMAGE1_PATH, "wb") as f:
        f.write(response.content)

def merge_with_logo():
    try:
        image1 = Image.open(IMAGE1_PATH).convert("RGB")
        image1 = image1.resize((600, 400))

        logo = Image.open(LOGO_PATH).convert("RGBA")
        logo = logo.rotate(45, expand=True)
        logo = logo.resize((100, 100))

        image1.paste(logo, (20, 20), logo)
        image1.save(MERGED_IMAGE_PATH)
    except FileNotFoundError:
        print("Logo non trouvé. Image utilisée sans logo.")
        image1 = Image.open(IMAGE1_PATH).resize((600, 400))
        image1.save(MERGED_IMAGE_PATH)

def create_word_doc(title, author, word_lengths):
    doc = Document()
    doc.add_heading(title, 0).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph(f"Auteur du livre : {author}", style="Normal")
    doc.add_paragraph(f"Auteur du rapport : {AUTEUR_RAPPORT}", style="Normal")
    doc.add_picture(MERGED_IMAGE_PATH, width=Inches(4))
    doc.add_page_break()

    doc.add_heading("Analyse du premier chapitre", level=1)
    doc.add_picture(GRAPH_PATH, width=Inches(5))
    doc.add_paragraph("Résumé de l’analyse :", style="Intense Quote")
    doc.add_paragraph(
        f"- Nombre de paragraphes : {len(word_lengths)}\n"
        f"- Nombre total de mots : {sum(word_lengths)}\n"
        f"- Nombre minimal de mots dans un paragraphe : {min(word_lengths)}\n"
        f"- Nombre maximal de mots dans un paragraphe : {max(word_lengths)}\n"
        f"- Moyenne de mots par paragraphe : {round(sum(word_lengths)/len(word_lengths), 2)}\n"
        f"- Source : Project Gutenberg ({BOOK_URL})"
    )
    doc.save(DOCX_PATH)

def main():
    print("Téléchargement du livre...")
    book_text = download_book()

    print("Extraction des informations...")
    title, author, paragraphs = extract_info(book_text)

    print("Analyse des paragraphes...")
    lengths = analyze_paragraphs(paragraphs)

    print("Téléchargement de l'image principale...")
    download_image()

    print("Fusion avec le logo (si présent)...")
    merge_with_logo()

    print("Création du document Word...")
    create_word_doc(title, author, lengths)

    print(f"✅ Rapport généré : {DOCX_PATH}")
